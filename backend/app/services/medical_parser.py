from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
import asyncio
import json
import os
import re
from transformers import AutoTokenizer
import spacy

# AI和NLP工具
# NLTK数据下载检查
# 文档处理库
from io import BytesIO
from nltk.tokenize import sent_tokenize, word_tokenize
import PyPDF2
import docx
import hashlib
import nltk
import pdfplumber

from app.utils.mock_spacy import spacy
from app.core.config import get_settings
from app.core.llm_client import get_llm_client
from app.core.logger import get_logger

"""
智能医学文档解析器
Hierarchical Medical Document Parser
节点1：智能文档解析层
"""

try:
    SPACY_AVAILABLE = True
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    TRANSFORMERS_AVAILABLE = False

# 创建模拟的AutoTokenizer
class MockAutoTokenizer:
    @staticmethod
    def from_pretrained(model_name):
        return MockTokenizer()

class MockTokenizer:
    def __call__(self, text, **kwargs):
        words = text.split()
        return {"input_ids": list(range(len(words))), "attention_mask": [1] * len(words)}

AutoTokenizer = MockAutoTokenizer

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# 可选的NLTK组件（仅在需要时下载）
# from nltk.chunk import ne_chunk
# from nltk.tag import pos_tag
# from nltk.tree import Tree

# 内部模块
logger = get_logger(__name__)

# 导入共享的文件类型常量
from shared.constants.file_types import FileType as DocumentType, FileValidationConfig

# 为了向后兼容，创建FileType别名
FileType = DocumentType


class EvidenceLevel(str, Enum):
    """证据等级"""
    A = "A"  # 高质量证据
    B = "B"  # 中等质量证据
    C = "C"  # 低质量证据
    D = "D"  # 专家意见
    NA = "NA"  # 不适用


class SectionType(str, Enum):
    """章节类型"""
    TITLE = "title"
    ABSTRACT = "abstract"
    INTRODUCTION = "introduction"
    METHODS = "methods"
    RECOMMENDATIONS = "recommendations"
    DISCUSSION = "discussion"
    CONCLUSION = "conclusion"
    REFERENCES = "references"
    APPENDIX = "appendix"
    TABLES = "tables"
    FIGURES = "figures"
    ALGORITHM = "algorithm"
    EVIDENCE = "evidence"
    OTHER = "other"


class ContentType(str, Enum):
    """内容类型"""
    TEXT = "text"
    TABLE = "table"
    ALGORITHM = "algorithm"
    LIST = "list"
    QUOTE = "quote"
    CODE = "code"
    FOOTNOTE = "footnote"
    REFERENCE = "reference"


@dataclass
class DocumentMetadata:
    """文档元数据"""
    file_path: str
    file_hash: str
    file_size: int
    file_type: DocumentType
    page_count: Optional[int] = None
    word_count: int = 0
    char_count: int = 0
    creation_date: Optional[str] = None
    modification_date: Optional[str] = None
    authors: List[str] = field(default_factory=list)
    title: Optional[str] = None
    abstract: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    doi: Optional[str] = None
    journal: Optional[str] = None
    publication_year: Optional[int] = None
    language: str = "zh-CN"


@dataclass
class MedicalEntity:
    """医学实体"""
    text: str
    label: str  # DISEASE, SYMPTOM, DRUG, TREATMENT, TEST等
    start: int
    end: int
    confidence: float
    ontology_id: Optional[str] = None  # 医学术语本体ID
    synonyms: List[str] = field(default_factory=list)


@dataclass
class DocumentSection:
    """文档章节"""
    section_id: str
    title: str
    section_type: SectionType
    content: str
    start_position: int
    end_position: int
    level: int  # 章节层级
    parent_id: Optional[str] = None
    children_ids: List[str] = field(default_factory=list)
    entities: List[MedicalEntity] = field(default_factory=list)
    evidence_level: Optional[EvidenceLevel] = None
    keywords: List[str] = field(default_factory=list)
    summary: Optional[str] = None


@dataclass
class MedicalTable:
    """医学表格"""
    table_id: str
    title: Optional[str]
    headers: List[str]
    rows: List[List[str]]
    caption: Optional[str]
    content_text: str  # 表格内容的文本表示
    metadata: Dict[str, Any] = field(default_factory=dict)
    evidence_level: Optional[EvidenceLevel] = None
    interpretation: Optional[str] = None  # AI解读


@dataclass
class ClinicalAlgorithm:
    """临床算法"""
    algorithm_id: str
    title: str
    steps: List[Dict[str, Any]]  # 算法步骤
    decision_points: List[Dict[str, Any]]  # 决策点
    flowchart_text: str
    source_section: str
    evidence_level: Optional[EvidenceLevel] = None
    target_population: Optional[str] = None


@dataclass
class EvidenceReference:
    """证据引用"""
    ref_id: str
    citation_text: str
    evidence_level: EvidenceLevel
    study_type: Optional[str]  # RCT, meta-analysis, cohort等
    sample_size: Optional[int]
    effect_size: Optional[float]
    confidence_interval: Optional[str]
    p_value: Optional[float]
    year: Optional[int]
    authors: List[str] = field(default_factory=list)
    journal: Optional[str] = None


@dataclass
class DocumentChunk:
    """文档分块"""
    chunk_id: str
    content: str
    chunk_type: ContentType
    section_id: Optional[str]
    start_position: int
    end_position: int
    word_count: int
    entities: List[MedicalEntity] = field(default_factory=list)
    keywords: List[str] = field(default_factory=list)
    summary: Optional[str] = None
    references: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    semantic_boundary: bool = False  # 是否为语义边界


@dataclass
class EvidenceHierarchy:
    """证据等级体系"""
    guideline_basis: str  # 指南制定依据
    grading_system: str  # 证据分级系统
    primary_evidence: List[EvidenceReference] = field(default_factory=list)
    secondary_evidence: List[EvidenceReference] = field(default_factory=list)
    expert_opinion: List[EvidenceReference] = field(default_factory=list)
    last_updated: Optional[str] = None


@dataclass
class MedicalDocument:
    """医学文档"""
    document_id: str
    metadata: DocumentMetadata
    sections: List[DocumentSection]
    tables: List[MedicalTable]
    algorithms: List[ClinicalAlgorithm]
    evidence_hierarchy: EvidenceHierarchy
    chunks: List[DocumentChunk] = field(default_factory=list)
    full_text: str = ""
    language_detected: str = "zh-CN"
    processing_notes: List[str] = field(default_factory=list)


class HierarchicalMedicalParser:
    """智能医学文档解析器"""

    def __init__(self):
        self.settings = get_settings()
        self.llm_client = None
        self._spacy_model = None
        self._tokenizer = None

        # 医学关键词模式
        self.medical_patterns = {
            'evidence_level': re.compile(
                r'(证据等级|证据级别|循证等级|Grade\s+[A-D]|Level\s+[1-5])',
                re.IGNORECASE
            ),
            'recommendation': re.compile(
                r'(推荐|建议|应该|should|recommend|guideline)',
                re.IGNORECASE
            ),
            'clinical_question': re.compile(
                r'(临床问题|研究问题|PICO|Clinical\s+question)',
                re.IGNORECASE
            ),
            'algorithm': re.compile(
                r'(流程图|算法|诊断流程|治疗流程|algorithm|flowchart)',
                re.IGNORECASE
            ),
            'adverse_effect': re.compile(
                r'(不良反应|副作用|不良事件|adverse\s+effect|side\s+effect)',
                re.IGNORECASE
            ),
            'contraindication': re.compile(
                r'(禁忌症|禁忌|contraindication)',
                re.IGNORECASE
            )
        }

    async def _get_llm_client(self):
        """获取LLM客户端"""
        if self.llm_client is None:
            self.llm_client = await get_llm_client()
        return self.llm_client

    async def _get_spacy_model(self):
        """获取Spacy模型"""
        if self._spacy_model is None:
            try:
                # 尝试加载中文医学模型
                self._spacy_model = spacy.load("zh_core_web_sm")
                logger.info("成功加载中文Spacy模型")
            except OSError:
                try:
                    # 如果没有中文模型，使用英文模型
                    self._spacy_model = spacy.load("en_core_web_sm")
                    logger.warning("中文Spacy模型未找到，使用英文模型")
                except OSError:
                    logger.error("Spacy模型加载失败，某些NLP功能将受限")
                    self._spacy_model = None
        return self._spacy_model

    async def _get_tokenizer(self):
        """获取分词器"""
        if self._tokenizer is None:
            try:
                self._tokenizer = AutoTokenizer.from_pretrained("bert-base-chinese")
            except Exception:
                # 如果无法加载，使用简单分词
                self._tokenizer = None
        return self._tokenizer

    def _calculate_file_hash(self, file_content: bytes) -> str:
        """计算文件哈希"""
        return hashlib.sha256(file_content).hexdigest()

    def _validate_file_size(self, file_size: int) -> None:
        """验证文件大小（使用共享常量）"""
        if file_size > FileValidationConfig.MAX_SIZE:
            raise ValueError(f"文件过大，超过限制 {FileValidationConfig.MAX_SIZE_MB}MB")

    def _validate_file_path(self, file_path: Path) -> None:
        """验证文件路径安全性"""
        # 转换为绝对路径并解析符号链接
        try:
            abs_path = file_path.resolve()
        except Exception as e:
            raise ValueError(f"无效的文件路径: {e}")

        # 检查文件是否存在
        if not abs_path.exists():
            raise FileNotFoundError(f"文件不存在: {abs_path}")

        # 检查是否为文件（不是目录）
        if not abs_path.is_file():
            raise ValueError(f"路径不是文件: {abs_path}")

        # 检查文件扩展名（使用共享常量）
        from shared.constants.file_types import ALLOWED_FILE_EXTENSIONS
        if abs_path.suffix.lower() not in ALLOWED_FILE_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {abs_path.suffix}")

        return abs_path

    async def parse_medical_document(self, file_path: str) -> MedicalDocument:
        """解析医学文档，返回结构化的医学文档对象"""
        logger.info(f"开始解析医学文档: {file_path}")

        # 1. 验证文件路径和大小
        file_path = Path(file_path)
        validated_path = self._validate_file_path(file_path)

        # 2. 读取文件并计算元数据
        try:
            file_content = validated_path.read_bytes()
            self._validate_file_size(len(file_content))
            file_hash = self._calculate_file_hash(file_content)
        except PermissionError:
            raise ValueError(f"文件读取权限不足: {validated_path}")
        except Exception as e:
            raise ValueError(f"文件读取失败: {e}")

        # 确定文档类型
        try:
            file_ext = validated_path.suffix.lower().lstrip('.')
            file_type = DocumentType(file_ext)
        except ValueError:
            raise ValueError(f"不支持的文档类型: {validated_path.suffix}")

        # 创建元数据
        metadata = DocumentMetadata(
            file_path=str(validated_path),
            file_hash=file_hash,
            file_size=len(file_content),
            file_type=file_type
        )

        # 2. 根据文档类型提取文本
        full_text, extracted_metadata = await self._extract_text_by_type(file_path, file_type)

        # 更新元数据
        metadata.word_count = len(full_text.split())
        metadata.char_count = len(full_text)
        metadata.language = "zh-CN"  # 假设为中文

        if extracted_metadata:
            for key, value in extracted_metadata.items():
                if hasattr(metadata, key):
                    setattr(metadata, key, value)

        # 3. 结构分析 - 使用LLM分析文档结构
        structure_analysis = await self._analyze_document_structure(full_text)

        # 4. 提取语义章节
        sections = self._extract_semantic_sections(full_text, structure_analysis)

        # 5. 提取表格
        tables = await self._extract_and_parse_tables(file_path, file_type, full_text)

        # 6. 提取临床算法
        algorithms = await self._extract_clinical_algorithms(full_text, sections)

        # 7. 提取证据等级体系
        evidence_hierarchy = await self._extract_evidence_hierarchy(full_text, sections)

        # 8. 创建医学文档对象
        document = MedicalDocument(
            document_id=f"doc_{file_hash[:16]}",
            metadata=metadata,
            sections=sections,
            tables=tables,
            algorithms=algorithms,
            evidence_hierarchy=evidence_hierarchy,
            full_text=full_text
        )

        logger.info(f"文档解析完成: {len(sections)}个章节, {len(tables)}个表格, {len(algorithms)}个算法")
        return document

    async def _extract_text_by_type(self, file_path: Path, file_type: DocumentType) -> Tuple[str, Dict[str, Any]]:
        """根据文档类型提取文本"""
        extracted_metadata = {}

        if file_type == DocumentType.PDF:
            text, meta = await self._extract_pdf_text(file_path)
            extracted_metadata.update(meta)
        elif file_type == DocumentType.DOCX:
            text, meta = await self._extract_docx_text(file_path)
            extracted_metadata.update(meta)
        elif file_type == DocumentType.TXT:
            text, meta = await self._extract_txt_text(file_path)
            extracted_metadata.update(meta)
        else:
            raise ValueError(f"不支持的文档类型: {file_type}")

        return text, extracted_metadata

    async def _extract_pdf_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """提取PDF文本"""
        text = ""
        metadata = {}

        try:
            # 使用pdfplumber提取文本（更好地处理表格和布局）
            with pdfplumber.open(file_path) as pdf:
                metadata['page_count'] = len(pdf.pages)
                metadata['authors'] = pdf.metadata.get('Author', '').split(';') if pdf.metadata.get('Author') else []
                metadata['title'] = pdf.metadata.get('Title', '')
                metadata['creation_date'] = pdf.metadata.get('CreationDate', '')

                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"

        except Exception as e:
            logger.warning(f"pdfplumber提取失败，尝试PyPDF2: {e}")

            # 备用方案：使用PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['page_count'] = len(pdf_reader.pages)

                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['authors'] = pdf_reader.metadata.get(
                        '/Author',
                        ''
                    ).split(';') if pdf_reader.metadata.get('/Author') else []

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"

        return text.strip(), metadata

    async def _extract_docx_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """提取DOCX文本"""
        doc = docx.Document(file_path)
        text = ""
        metadata = {}

        # 提取段落文本
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # 提取表格文本
        for table in doc.tables:
            text += "\n[表格]\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"

        # 提取文档属性
        core_props = doc.core_properties
        metadata['title'] = core_props.title or ""
        metadata['authors'] = [core_props.author] if core_props.author else []
        metadata['creation_date'] = str(core_props.created) if core_props.created else ""
        metadata['keywords'] = core_props.keywords.split(',') if core_props.keywords else []

        return text.strip(), metadata

    async def _extract_txt_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """提取TXT文本"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        # 尝试从文本开头提取元数据
        lines = text.split('\n')
        metadata = {}

        # 简单的标题提取（假设第一行是标题）
        if lines and len(lines[0]) < 200:  # 假设标题不会太长
            metadata['title'] = lines[0].strip()

        return text.strip(), metadata


# 全局解析器实例
medical_parser = HierarchicalMedicalParser()


async def get_medical_parser() -> HierarchicalMedicalParser:
    """获取医学文档解析器实例"""
    return medical_parser


# 便捷函数
async def parse_medical_document(file_path: str) -> MedicalDocument:
    """解析医学文档的便捷函数"""
    parser = await get_medical_parser()
    return await parser.parse_medical_document(file_path)


async def create_document_chunks(document: MedicalDocument, max_chunk_size: int = 800) -> List[DocumentChunk]:
    """创建文档分块的便捷函数"""
    parser = await get_medical_parser()
    return await parser.adaptive_chunking(document, max_chunk_size)